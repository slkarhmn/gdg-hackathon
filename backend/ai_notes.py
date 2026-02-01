# ai_notes.py

import os
import json
from openai import OpenAI  # Updated import
import tempfile
from docx import Document
from fpdf import FPDF
from flask import send_file, jsonify
import logging

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | ai_notes | %(message)s"
)

logger = logging.getLogger(__name__)


# Initialize OpenAI client (matching ai_chat.py pattern)
client = None

def initialize_openai():
    """Initialize OpenAI client with API key from environment"""
    global client
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY not found in environment variables")
    client = OpenAI(api_key=api_key)


def extract_text_from_note(note):
    """
    Extracts lecture note text from stored content (expects Note SQLAlchemy object).
    Supports content with 'text', 'body', or 'title' keys (frontend uses {title, body}).
    """
    try:
        logger.info(f"Extracting text from note ID={note.id}")
        content_dict = json.loads(note.content) if isinstance(note.content, str) else (note.content or {})
        text = content_dict.get("text", "")
        if not text:
            body = content_dict.get("body", "")
            title = content_dict.get("title", "")
            if isinstance(body, str) and isinstance(title, str):
                text = (f"# {title}\n\n" if title else "") + body
            elif isinstance(body, str):
                text = body
            elif isinstance(title, str):
                text = title

        logger.info(f"Extracted text length: {len(text)} characters")
        return text or ""
    except Exception as e:
        logger.exception("Failed to extract text from note")
        raise

def generate_summary_and_questions(text):
    """
    Calls OpenAI to generate summary notes and exam-style questions.
    """
    # Initialize client if not already done
    if not client:
        initialize_openai()
    
    print(text)
    if not text.strip():
        logger.warning("Empty lecture text received for AI generation")
        return "", ""

    logger.info("Starting OpenAI summary + questions generation")

    prompt = (
        f"Given the following lecture note, do two things:\n"
        f"1. Write a concise, student-friendly summary with clear headings and bullet points.\n"
        f"2. Write 3-5 exam-style questions that could come from this content.\n\n"
        f"Lecture Note:\n{text}\n"
        f"---\n"
        f"Your response format:\n"
        f"## Summary\n[summary here]\n\n## Possible Exam Questions\n- Q1\n- Q2\n- Q3"
    )

    try:
        logger.info("Calling OpenAI ChatCompletion API")

        # Updated to use new API syntax
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=800,
            temperature=0.4,
        )

        logger.info("OpenAI response received successfully")

        # Updated to access response content
        result = response.choices[0].message.content

    except Exception as e:
        logger.exception("OpenAI API call FAILED")
        raise

    try:
        if "## Possible Exam Questions" in result:
            summary, questions = result.split("## Possible Exam Questions", 1)
            summary = summary.replace("## Summary", "").strip()
            questions = questions.strip()
        else:
            summary = result
            questions = ""

        logger.info("Successfully parsed AI output")
        return summary, questions

    except Exception as e:
        logger.exception("Failed to parse OpenAI output")
        raise



def export_note_as_docx(note, summary, questions):
    try:
        logger.info(f"Exporting note ID={note.id} as DOCX")

        doc = Document()
        doc.add_heading(note.subject or "Lecture Notes", level=1)
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(summary)

        doc.add_heading("Possible Exam Questions", level=2)
        for line in questions.split("\n"):
            doc.add_paragraph(line, style='ListBullet')

        temp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc.save(temp.name)
        temp.seek(0)

        logger.info("DOCX export successful")
        return temp

    except Exception:
        logger.exception("DOCX export failed")
        raise



def export_note_as_pdf(note, summary, questions):
    try:
        logger.info(f"Exporting note ID={note.id} as PDF")

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=14)
        pdf.multi_cell(0, 10, note.subject or "Lecture Notes")

        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, "Summary\n" + summary)
        pdf.ln(5)
        pdf.multi_cell(0, 10, "Possible Exam Questions\n" + questions)

        temp = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        pdf.output(temp.name)
        temp.seek(0)

        logger.info("PDF export successful")
        return temp

    except Exception:
        logger.exception("PDF export failed")
        raise